import os
import re
import json
import datetime
import io
import logging
import uuid
from flask import Flask, request, render_template, jsonify, session, redirect, url_for, send_file, flash
import google.generativeai as genai
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
import docx
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile

# Configurar logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['RESULT_FOLDER'] = 'results'  # Pasta para guardar resultados temporários
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max upload size
app.config['SECRET_KEY'] = '208d68f338ce335f60117b11b4072a32'  # Chave fixa para sessões

# Create uploads folder if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['RESULT_FOLDER'], exist_ok=True)

# Configure Gemini API (you'll need to set your API key in environment variables)
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

# Configure Gemini
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    logger.info("Gemini API configurada com sucesso.")
else:
    logger.warning("GEMINI_API_KEY environment variable not set")

# Unified prompt for Gemini
PROMPT = """
Você é um assistente jurídico especializado em extração de dados e formatação de documentos jurídicos. Receberá dois documentos de texto extraídos de PDFs:

1. Uma **petição inicial** contendo os dados do autor, réu, fatos, fundamentos jurídicos e pedidos.
2. Um **modelo de contestação jurídica**, usado como referência de estrutura, estilo e linguagem.

Sua tarefa é:

### ETAPA 1: Extrair informações detalhadas da petição inicial no formato JSON:
```json
{
  "processo": {
    "numero": "",
    "comarca": "",
    "vara": "",
    "foro": ""
  },
  "autor": {
    "nome": "",
    "cpf_cnpj": "",
    "qualificacao": "",
    "endereco": "",
    "representacao": {
      "advogado": "",
      "oab": ""
    }
  },
  "reu": {
    "nome": "",
    "cnpj": "",
    "qualificacao": "",
    "endereco": "",
    "representacao": {
      "advogado": "",
      "oab": ""
    }
  },
  "objeto": "",
  "fatos": [
    {
      "numero": "",
      "descricao": "",
      "data": "",
      "valor": ""
    }
  ],
  "fundamentos": [
    {
      "tipo": "",
      "descricao": "",
      "artigos": []
    }
  ],
  "pedidos": [
    {
      "numero": "",
      "descricao": "",
      "valor": ""
    }
  ],
  "documentos": [
    {
      "tipo": "",
      "descricao": ""
    }
  ]
}
```

### ETAPA 2: Gerar uma contestação jurídica seguindo EXATAMENTE a estrutura formal do CPC:

1. **ENDEREÇAMENTO**
   ```
   AO JUÍZO DA [NÚMERO]ª VARA CÍVEL DO FORO [NOME] DA COMARCA DE [CIDADE] ([UF])
   ```

2. **QUALIFICAÇÃO DAS PARTES**
   ```
   [NOME DA PARTE RÉ], pessoa jurídica de direito privado, inscrita no CNPJ sob o n.º [CNPJ], com sede na [ENDEREÇO], neste ato representada por seu advogado, instrumento de mandato com poderes para o foro em anexo, com escritório profissional sito à [ENDEREÇO DO ESCRITÓRIO], onde recebe intimações, nos termos do artigo 319, inciso I, do CPC.
   ```

3. **NÚMERO DO PROCESSO**
   ```
   Processo n.º [NÚMERO]
   ```

4. **TÍTULO**
   ```
   CONTESTAÇÃO
   ```

5. **PRELIMINARMENTE**
   - Argumentos processuais numerados
   - Cada argumento com fundamentação legal
   - Citações de artigos do CPC

6. **DO MÉRITO**
   - Fatos e fundamentos jurídicos
   - Argumentação detalhada
   - Jurisprudência relevante
   - Artigos de lei aplicáveis

7. **DOS PEDIDOS**
   - Pedidos numerados
   - Fundamentação de cada pedido
   - Valores e prazos quando aplicável

8. **DOCUMENTOS ANEXOS**
   ```
   DOCUMENTOS ANEXOS:
   1. Procuração
   2. [Outros documentos]
   ```

9. **CONCLUSÃO E ASSINATURA**
   ```
   Termos em que,
   Pede deferimento.

   [CIDADE], [DATA].

   [NOME DO ADVOGADO]
   OAB/[UF] [NÚMERO]
   ```

⚠️ REGRAS DE FORMATAÇÃO:
1. Usar fonte Times New Roman, tamanho 12
2. Espaçamento entre linhas de 1,5
3. Margens: superior e esquerda 3cm, inferior e direita 2cm
4. Alinhamento justificado
5. Recuo de 2cm para início de parágrafo
6. Numeração de páginas no canto inferior direito
7. Seções em negrito e maiúsculas
8. Artigos de lei em negrito
9. Jurisprudência em itálico

### SAÍDA:

1. Primeiro, retorne o JSON estruturado com todos os dados extraídos
2. Em seguida, retorne a contestação formatada seguindo EXATAMENTE a estrutura acima
"""

def save_result_to_file(result):
    """Salvar resultado em arquivo temporário e retornar o ID"""
    result_id = str(uuid.uuid4())
    result_path = os.path.join(app.config['RESULT_FOLDER'], f"{result_id}.txt")
    
    try:
        with open(result_path, 'w', encoding='utf-8') as f:
            f.write(result)
        logger.info(f"Resultado salvo em arquivo: {result_path}")
        return result_id
    except Exception as e:
        logger.error(f"Erro ao salvar resultado em arquivo: {str(e)}")
        return None

def get_result_from_file(result_id):
    """Recuperar resultado do arquivo temporário"""
    if not result_id:
        return None
        
    result_path = os.path.join(app.config['RESULT_FOLDER'], f"{result_id}.txt")
    
    try:
        if not os.path.exists(result_path):
            logger.error(f"Arquivo de resultado não encontrado: {result_path}")
            return None
            
        with open(result_path, 'r', encoding='utf-8') as f:
            result = f.read()
        logger.info(f"Resultado recuperado do arquivo: {result_path}")
        return result
    except Exception as e:
        logger.error(f"Erro ao ler resultado do arquivo: {str(e)}")
        return None

def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using PyMuPDF"""
    text = ""
    try:
        # Verificar se o arquivo existe antes de tentar abri-lo
        if not os.path.exists(pdf_path):
            logger.error(f"Arquivo PDF não encontrado: {pdf_path}")
            return f"Error: O arquivo {pdf_path} não foi encontrado"
        
        # Open the PDF file
        pdf_document = fitz.open(pdf_path)
        
        # Get the number of pages
        num_pages = len(pdf_document)
        logger.info(f"Extraindo texto de PDF com {num_pages} páginas")
        
        # Extract text from each page
        for page_num in range(num_pages):
            page = pdf_document.load_page(page_num)
            text += page.get_text()
            
        # Close the PDF file
        pdf_document.close()
        
        logger.info(f"Texto extraído com sucesso: {len(text)} caracteres")
        return text
    except Exception as e:
        logger.error(f"Erro ao extrair texto do PDF: {str(e)}")
        return f"Error extracting text from PDF: {str(e)}"

def process_pdfs_with_gemini(peticao_pdf_path, modelo_pdf_path):
    try:
        # Extract text from PDFs using PyMuPDF
        peticao_text = extract_text_from_pdf(peticao_pdf_path)
        modelo_text = extract_text_from_pdf(modelo_pdf_path)
        
        # Verificar se o texto foi extraído corretamente
        if not peticao_text or peticao_text.startswith("Error"):
            logger.error(f"Erro na extração do texto da petição: {peticao_text}")
            return f"Erro ao extrair texto da petição inicial: {peticao_text}"
            
        if not modelo_text or modelo_text.startswith("Error"):
            logger.error(f"Erro na extração do texto do modelo: {modelo_text}")
            return f"Erro ao extrair texto do modelo de contestação: {modelo_text}"
        
        # Initialize Gemini model
        model = genai.GenerativeModel('gemini-2.0-flash')
        logger.info("Modelo Gemini inicializado. Enviando conteúdo para processamento...")
        
        # Prepare content for Gemini
        contents = [
            PROMPT,
            f"### PETIÇÃO INICIAL:\n{peticao_text}",
            f"### MODELO DE CONTESTAÇÃO:\n{modelo_text}"
        ]
        
        # Generate response
        response = model.generate_content(contents)
        
        # Verificar resposta
        if response and hasattr(response, 'text') and response.text:
            logger.info(f"Resposta recebida do Gemini: {len(response.text)} caracteres")
            return response.text
        else:
            logger.error("Resposta vazia ou inválida do Gemini")
            return "Erro: Resposta vazia ou inválida do Gemini. Verifique se sua API key está correta e tente novamente."
    except Exception as e:
        logger.error(f"Erro ao processar PDFs com Gemini: {str(e)}")
        return f"Erro ao processar PDFs: {str(e)}"

def extract_json_and_contestacao(response_text):
    """Extract JSON and contestação from Gemini response"""
    try:
        logger.info(f"Extraindo JSON e contestação do texto ({len(response_text)} caracteres)")
        
        # Try to find JSON using regex
        json_match = re.search(r'```json\s*(.*?)\s*```', response_text, re.DOTALL)
        
        if not json_match:
            # Try alternative pattern without code block markers
            json_match = re.search(r'({[\s\S]*?"pedidos"\s*:\s*\[[\s\S]*?\]\s*})', response_text)
        
        json_data = None
        if json_match:
            try:
                json_str = json_match.group(1)
                json_data = json.loads(json_str)
                # Remove the JSON part from response text to get contestação
                contestacao = response_text[json_match.end():].strip()
                logger.info(f"JSON extraído com sucesso. Contestação: {len(contestacao)} caracteres")
            except json.JSONDecodeError as e:
                logger.error(f"Erro ao decodificar JSON: {str(e)}")
                json_data = {"error": "JSON inválido no resultado"}
                contestacao = response_text
        else:
            logger.warning("JSON não encontrado no texto da resposta")
            json_data = {"error": "JSON não encontrado no resultado"}
            contestacao = response_text
        
        return json_data, contestacao
    except Exception as e:
        logger.error(f"Erro ao extrair JSON e contestação: {str(e)}")
        return {"error": f"Erro ao extrair dados: {str(e)}"}, response_text

def parse_contestacao_sections(text):
    """Parse contestação text into sections with hierarchical structure"""
    try:
        logger.info(f"Dividindo contestação em seções hierárquicas ({len(text)} caracteres)")
        
        # Padrões de seções principais
        main_sections = [
            r'PRELIMINARMENTE|PRELIMINAR',
            r'DO MÉRITO|MÉRITO',
            r'DOS PEDIDOS|DOS REQUERIMENTOS',
            r'DOCUMENTOS ANEXOS'
        ]
        
        # Padrões de subseções
        subsections = [
            r'(\d+)[\.\)]\s*([A-Z][^\.]+)',
            r'([a-z]\))\s*([A-Z][^\.]+)',
            r'(\d+)[\.\)]\s*([A-Z][^\.]+)'
        ]
        
        sections = []
        current_section = None
        current_subsection = None
        
        # Dividir o texto em linhas e remover duplicações
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        unique_lines = []
        for line in lines:
            if line not in unique_lines:
                unique_lines.append(line)
        
        for line in unique_lines:
            # Verificar se é uma seção principal
            for pattern in main_sections:
                if re.match(pattern, line, re.IGNORECASE):
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        "title": line.upper(),
                        "content": "",
                        "subsections": []
                    }
                    break
            
            # Verificar se é uma subseção
            if current_section:
                for pattern in subsections:
                    match = re.match(pattern, line)
                    if match:
                        if current_subsection:
                            current_section["subsections"].append(current_subsection)
                        current_subsection = {
                            "number": match.group(1),
                            "title": match.group(2).strip(),
                            "content": ""
                        }
                        break
                    else:
                        if current_subsection:
                            current_subsection["content"] += line + "<br>"
                        else:
                            current_section["content"] += line + "<br>"
        
        # Adicionar última subseção e seção
        if current_subsection:
            current_section["subsections"].append(current_subsection)
        if current_section:
            sections.append(current_section)
        
        # Se não encontrou seções, criar uma única seção
        if not sections:
            sections.append({
                "title": "CONTESTAÇÃO",
                "content": text.replace('\n', '<br>'),
                "subsections": []
            })
        
        logger.info(f"Contestação dividida em {len(sections)} seções principais")
        return sections
    except Exception as e:
        logger.error(f"Erro ao dividir contestação em seções: {str(e)}")
        return [{
            "title": "CONTESTAÇÃO",
            "content": text.replace('\n', '<br>'),
            "subsections": []
        }]

def create_word_document(contestacao_data):
    try:
        doc = docx.Document()
        
        # Configurar margens (em centímetros)
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(3)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)
        
        # Adicionar cabeçalho
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run("MINHA HONRA JUS\nExcelência em Documentos Jurídicos")
        header_run.bold = True
        header_run.font.size = Pt(14)
        
        # Adicionar endereçamento
        doc.add_paragraph()
        endereco = doc.add_paragraph()
        endereco.alignment = WD_ALIGN_PARAGRAPH.CENTER
        endereco.add_run("EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO").bold = True
        endereco.add_run(f"\nDA VARA CÍVEL DO FORO {contestacao_data['foro']} DA COMARCA DE {contestacao_data['comarca']}")
        
        # Adicionar informações do processo
        doc.add_paragraph()
        processo = doc.add_paragraph()
        processo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        processo.add_run(f"Processo n.º: {contestacao_data['numero_processo']}\n")
        processo.add_run(f"Autor: {contestacao_data['autor_nome']}\n")
        processo.add_run(f"Réu: {contestacao_data['reu_nome']}")
        
        # Adicionar título da contestação
        doc.add_paragraph()
        titulo = doc.add_paragraph()
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        titulo_run = titulo.add_run("CONTESTAÇÃO")
        titulo_run.bold = True
        titulo_run.font.size = Pt(14)
        
        # Adicionar seções
        for section in contestacao_data['secoes']:
            # Adicionar título da seção
            doc.add_paragraph()
            section_title = doc.add_paragraph()
            section_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            section_title_run = section_title.add_run(section['titulo'])
            section_title_run.bold = True
            section_title_run.font.size = Pt(12)
            
            # Adicionar parágrafos da seção
            for paragraph in section['paragrafos']:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(2)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.add_run(paragraph)
        
        # Adicionar documentos anexos
        doc.add_paragraph()
        docs_title = doc.add_paragraph()
        docs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        docs_title_run = docs_title.add_run("DOCUMENTOS ANEXOS")
        docs_title_run.bold = True
        docs_title_run.font.size = Pt(12)
        
        docs_list = [
            "Procuração com poderes especiais",
            "Documentos constitutivos da empresa ré",
            "Termos e Condições da plataforma Brazino777",
            "Registros de apostas do Autor",
            "Laudo pericial (quando disponível)"
        ]
        
        for doc_item in docs_list:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(2)
            p.add_run(f"• {doc_item}")
        
        # Adicionar assinatura
        doc.add_paragraph()
        assinatura = doc.add_paragraph()
        assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
        assinatura.add_run("Termos em que,\nPede deferimento.\n\n")
        assinatura.add_run(f"{datetime.datetime.now().strftime('%d/%m/%Y')}\n\n")
        assinatura.add_run("_____________________________\n")
        assinatura.add_run(f"{contestacao_data['advogado_nome']}\n")
        assinatura.add_run(f"OAB/{contestacao_data['advogado_estado']} {contestacao_data['advogado_numero']}")
        
        return doc
    except Exception as e:
        logger.error(f"Erro ao criar documento Word: {str(e)}")
        raise

def create_txt_document(contestacao_data):
    try:
        lines = []
        
        # Adicionar cabeçalho
        lines.append("MINHA HONRA JUS")
        lines.append("Excelência em Documentos Jurídicos")
        lines.append("")
        
        # Adicionar endereçamento
        lines.append("EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO")
        lines.append(f"DA VARA CÍVEL DO FORO {contestacao_data['foro']} DA COMARCA DE {contestacao_data['comarca']}")
        lines.append("")
        
        # Adicionar informações do processo
        lines.append(f"Processo n.º: {contestacao_data['numero_processo']}")
        lines.append(f"Autor: {contestacao_data['autor_nome']}")
        lines.append(f"Réu: {contestacao_data['reu_nome']}")
        lines.append("")
        
        # Adicionar título da contestação
        lines.append("CONTESTAÇÃO")
        lines.append("")
        
        # Adicionar seções
        for section in contestacao_data['secoes']:
            lines.append(section['titulo'])
            lines.append("")
            
            for paragraph in section['paragrafos']:
                lines.append("    " + paragraph)
                lines.append("")
        
        # Adicionar documentos anexos
        lines.append("DOCUMENTOS ANEXOS")
        lines.append("")
        
        docs_list = [
            "Procuração com poderes especiais",
            "Documentos constitutivos da empresa ré",
            "Termos e Condições da plataforma Brazino777",
            "Registros de apostas do Autor",
            "Laudo pericial (quando disponível)"
        ]
        
        for doc_item in docs_list:
            lines.append(f"• {doc_item}")
        
        lines.append("")
        lines.append("Termos em que,")
        lines.append("Pede deferimento.")
        lines.append("")
        lines.append(f"{datetime.datetime.now().strftime('%d/%m/%Y')}")
        lines.append("")
        lines.append("_____________________________")
        lines.append(f"{contestacao_data['advogado_nome']}")
        lines.append(f"OAB/{contestacao_data['advogado_estado']} {contestacao_data['advogado_numero']}")
        
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"Erro ao criar documento TXT: {str(e)}")
        raise

@app.route('/')
def index():
    logger.info("Página inicial acessada")
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    logger.info("Requisição POST recebida em /process")
    # Check if API key is configured
    if not GEMINI_API_KEY:
        logger.error("API key não configurada")
        return render_template('index.html', error='API key not configured. Set GEMINI_API_KEY environment variable.'), 500

    try:
        # Check if both files are present in the request
        if 'peticao' not in request.files or 'modelo' not in request.files:
            logger.error("Arquivos necessários não encontrados na requisição")
            return render_template('index.html', error='Ambos os arquivos (petição e modelo) são necessários'), 400
        
        peticao_file = request.files['peticao']
        modelo_file = request.files['modelo']
        
        # Check if files are empty
        if peticao_file.filename == '' or modelo_file.filename == '':
            logger.error("Arquivos vazios")
            return render_template('index.html', error='Nenhum arquivo selecionado'), 400
        
        # Verificar extensão dos arquivos
        if not peticao_file.filename.lower().endswith('.pdf') or not modelo_file.filename.lower().endswith('.pdf'):
            logger.error("Arquivos não são PDFs")
            return render_template('index.html', error='Os arquivos devem ser PDFs'), 400
        
        # Save files with unique filenames to avoid conflicts
        secure_peticao_filename = secure_filename(f"{uuid.uuid4()}_{peticao_file.filename}")
        secure_modelo_filename = secure_filename(f"{uuid.uuid4()}_{modelo_file.filename}")
        
        logger.info(f"Salvando arquivos: {secure_peticao_filename} e {secure_modelo_filename}")
        peticao_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_peticao_filename)
        modelo_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_modelo_filename)
        
        peticao_file.save(peticao_path)
        modelo_file.save(modelo_path)
        
        # Verificar se os arquivos foram salvos corretamente
        if not os.path.exists(peticao_path):
            logger.error(f"Falha ao salvar o arquivo da petição: {peticao_path}")
            return render_template('index.html', error='Falha ao salvar o arquivo da petição. Tente novamente.'), 500
        
        if not os.path.exists(modelo_path):
            logger.error(f"Falha ao salvar o arquivo do modelo: {modelo_path}")
            return render_template('index.html', error='Falha ao salvar o arquivo do modelo. Tente novamente.'), 500
        
        # Process with Gemini
        logger.info("Processando PDFs com Gemini")
        result = process_pdfs_with_gemini(peticao_path, modelo_path)
        
        # Verificar resultado
        if not result or result.startswith("Erro"):
            logger.error(f"Erro no processamento: {result}")
            # Clean up uploaded files
            if os.path.exists(peticao_path):
                os.remove(peticao_path)
            if os.path.exists(modelo_path):
                os.remove(modelo_path)
            return render_template('index.html', error=result), 500
        
        # Clean up uploaded files
        logger.info("Removendo arquivos temporários")
        try:
            os.remove(peticao_path)
            os.remove(modelo_path)
        except Exception as e:
            logger.warning(f"Erro ao remover arquivos temporários: {str(e)}")
        
        # Salvar resultado em arquivo em vez de usar sessão
        result_id = save_result_to_file(result)
        if not result_id:
            logger.error("Falha ao salvar resultado em arquivo")
            return render_template('index.html', error="Falha ao salvar resultado em arquivo"), 500
        
        # Guarda o ID do resultado na sessão como backup
        session['result_id'] = result_id
        
        # Redirect to formatted result page
        logger.info(f"Redirecionando para página de resultado com ID: {result_id}")
        return redirect(url_for('resultado', id=result_id))
    
    except Exception as e:
        logger.exception(f"Exceção não tratada: {str(e)}")
        # Tratar qualquer exceção não prevista
        return render_template('index.html', error=f'Erro ao processar: {str(e)}'), 500

@app.route('/resultado')
def resultado():
    logger.info("Página de resultado acessada")
    
    # Pegar ID do resultado da URL
    result_id = request.args.get('id')
    
    # Se não tem ID na URL, tentar pegar da sessão
    if not result_id:
        logger.warning("ID não encontrado na URL, tentando obter da sessão")
        result_id = session.get('result_id')
    
    if not result_id:
        logger.error("ID do resultado não encontrado")
        return render_template('index.html', error='Nenhum resultado encontrado. Por favor, envie os documentos novamente.'), 400
        
    # Recuperar o resultado do arquivo
    result = get_result_from_file(result_id)
    
    if not result:
        logger.error(f"Resultado não encontrado para o ID: {result_id}")
        return render_template('index.html', error='Resultado não encontrado. Por favor, envie os documentos novamente.'), 400
    
    try:
        # Extract JSON and contestação from the result
        logger.info("Extraindo JSON e contestação do resultado")
        json_data, contestacao_text = extract_json_and_contestacao(result)
        
        # Verificar se temos uma contestação
        if not contestacao_text or len(contestacao_text) < 50:
            logger.error(f"Contestação muito curta ou vazia: {contestacao_text}")
            return render_template('index.html', error='A contestação gerada está vazia ou inválida. Por favor, tente novamente.'), 400
        
        # Parse contestação into sections
        logger.info("Dividindo contestação em seções")
        contestacao_sections = parse_contestacao_sections(contestacao_text)
        
        # Get current date
        data_atual = datetime.datetime.now().strftime("%d/%m/%Y")
        
        # Extrair informações específicas do JSON
        autor_nome = json_data.get('autor', {}).get('nome', '')
        reu_nome = json_data.get('reu', {}).get('nome', '')
        
        # Return the rendered template
        logger.info("Renderizando template de resultado")
        return render_template('resultado.html', 
                              json_data=json.dumps(json_data, indent=2, ensure_ascii=False),
                              contestacao_sections=contestacao_sections,
                              data_atual=data_atual,
                              result_id=result_id,
                              autor_nome=autor_nome,
                              reu_nome=reu_nome,
                              comarca='São Paulo',  # Pode ser extraído do JSON se disponível
                              numero_processo='',   # Pode ser extraído do JSON se disponível
                              advogado_nome='GUILHERME KASCHNY BASTIAN',
                              advogado_estado='SP',
                              advogado_numero='266.795')
    except Exception as e:
        logger.exception(f"Erro ao renderizar página de resultado: {str(e)}")
        return render_template('index.html', error=f'Erro ao renderizar resultado: {str(e)}'), 500

@app.route('/download/docx')
def download_docx():
    try:
        # Obter dados da contestação dos argumentos da requisição
        contestacao_data = {
            'foro': request.args.get('foro', '[FORO]'),
            'comarca': request.args.get('comarca', 'SÃO PAULO'),
            'numero_processo': request.args.get('numero_processo', '[NÚMERO DO PROCESSO]'),
            'autor_nome': request.args.get('autor_nome', '[AUTOR]'),
            'reu_nome': request.args.get('reu_nome', '[RÉU]'),
            'advogado_nome': request.args.get('advogado_nome', '[NOME DO ADVOGADO]'),
            'advogado_estado': request.args.get('advogado_estado', 'XX'),
            'advogado_numero': request.args.get('advogado_numero', '000000'),
            'secoes': json.loads(request.args.get('secoes', '[]'))
        }
        
        # Criar documento Word
        doc = create_word_document(contestacao_data)
        
        # Salvar documento temporariamente
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        
        # Enviar arquivo para download
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'contestacao_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.error(f"Erro ao gerar DOCX: {str(e)}")
        return jsonify({'error': 'Erro ao gerar documento Word'}), 500

@app.route('/download/txt')
def download_txt():
    try:
        # Obter dados da contestação dos argumentos da requisição
        contestacao_data = {
            'foro': request.args.get('foro', '[FORO]'),
            'comarca': request.args.get('comarca', 'SÃO PAULO'),
            'numero_processo': request.args.get('numero_processo', '[NÚMERO DO PROCESSO]'),
            'autor_nome': request.args.get('autor_nome', '[AUTOR]'),
            'reu_nome': request.args.get('reu_nome', '[RÉU]'),
            'advogado_nome': request.args.get('advogado_nome', '[NOME DO ADVOGADO]'),
            'advogado_estado': request.args.get('advogado_estado', 'XX'),
            'advogado_numero': request.args.get('advogado_numero', '000000'),
            'secoes': json.loads(request.args.get('secoes', '[]'))
        }
        
        # Criar documento TXT
        content = create_txt_document(contestacao_data)
        
        # Salvar arquivo temporariamente
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt')
        with open(temp_file.name, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Enviar arquivo para download
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'contestacao_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.txt',
            mimetype='text/plain'
        )
    except Exception as e:
        logger.error(f"Erro ao gerar TXT: {str(e)}")
        return jsonify({'error': 'Erro ao gerar arquivo de texto'}), 500

@app.route('/api/process', methods=['POST'])
def api_process():
    """API endpoint for compatibility with previous implementation"""
    logger.info("Requisição para API recebida")
    # Check if API key is configured
    if not GEMINI_API_KEY:
        logger.error("API key não configurada")
        return jsonify({
            'error': 'API key not configured. Set GEMINI_API_KEY environment variable.'
        }), 500

    try:
        # Check if both files are present in the request
        if 'peticao' not in request.files or 'modelo' not in request.files:
            logger.error("Arquivos necessários não encontrados na requisição para API")
            return jsonify({
                'error': 'Ambos os arquivos (petição e modelo) são necessários'
            }), 400
        
        peticao_file = request.files['peticao']
        modelo_file = request.files['modelo']
        
        # Check if files are empty
        if peticao_file.filename == '' or modelo_file.filename == '':
            logger.error("Arquivos vazios para API")
            return jsonify({
                'error': 'Nenhum arquivo selecionado'
            }), 400
        
        # Save files with unique filenames to avoid conflicts
        secure_peticao_filename = secure_filename(f"{uuid.uuid4()}_{peticao_file.filename}")
        secure_modelo_filename = secure_filename(f"{uuid.uuid4()}_{modelo_file.filename}")
        
        logger.info(f"Salvando arquivos para API: {secure_peticao_filename} e {secure_modelo_filename}")
        peticao_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_peticao_filename)
        modelo_path = os.path.join(app.config['UPLOAD_FOLDER'], secure_modelo_filename)
        
        peticao_file.save(peticao_path)
        modelo_file.save(modelo_path)
        
        # Verificar se os arquivos foram salvos corretamente
        if not os.path.exists(peticao_path):
            logger.error(f"Falha ao salvar o arquivo da petição para API: {peticao_path}")
            return jsonify({
                'error': 'Falha ao salvar o arquivo da petição. Tente novamente.'
            }), 500
        
        if not os.path.exists(modelo_path):
            logger.error(f"Falha ao salvar o arquivo do modelo para API: {modelo_path}")
            return jsonify({
                'error': 'Falha ao salvar o arquivo do modelo. Tente novamente.'
            }), 500
        
        # Process with Gemini
        logger.info("Processando PDFs com Gemini para API")
        result = process_pdfs_with_gemini(peticao_path, modelo_path)
        
        # Clean up uploaded files
        logger.info("Removendo arquivos temporários para API")
        try:
            os.remove(peticao_path)
            os.remove(modelo_path)
        except Exception as e:
            logger.warning(f"Erro ao remover arquivos temporários para API: {str(e)}")
        
        # Para API, retornar JSON
        logger.info("Extraindo JSON e contestação para API")
        json_data, contestacao = extract_json_and_contestacao(result)
        
        # Salvar resultado em arquivo também para compatibilidade
        result_id = save_result_to_file(result)
        
        return jsonify({
            'result': result,
            'json_data': json_data,
            'contestacao': contestacao,
            'result_id': result_id
        })
    
    except Exception as e:
        logger.exception(f"Erro na API: {str(e)}")
        return jsonify({
            'error': f'Erro ao processar: {str(e)}'
        }), 500

@app.route('/debug/session_test')
def debug_session_test():
    """Rota para testar se a sessão está funcionando corretamente"""
    session['test_value'] = 'debug_ok_' + datetime.datetime.now().strftime('%H%M%S')
    sample_text = """
    EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA VARA CÍVEL DA COMARCA DE SÃO PAULO - SP
    
    DOS FATOS
    
    Trata-se de ação de cobrança proposta pelo autor, alegando ser credor do réu.
    
    DO DIREITO
    
    Não procedem as alegações do autor, pois os valores já foram pagos.
    
    DOS PEDIDOS
    
    Ante o exposto, requer-se a improcedência dos pedidos do autor.
    """
    
    # Simular resposta do Gemini com JSON e contestação
    fake_gemini_response = """```json
    {
      "autor": {
        "nome": "João da Silva",
        "cpf_cnpj": "123.456.789-00"
      },
      "reu": {
        "nome": "Empresa ABC Ltda",
        "cpf_cnpj": "12.345.678/0001-00"
      },
      "objeto": "Cobrança de valores",
      "fatos": [
        "O autor alega ser credor do réu"
      ],
      "fundamentos": [
        "Artigo 397 do Código Civil"
      ],
      "pedidos": [
        "Pagamento de R$ 10.000,00"
      ]
    }
    ```
    
    EXCELENTÍSSIMO SENHOR DOUTOR JUIZ DE DIREITO DA VARA CÍVEL DA COMARCA DE SÃO PAULO - SP
    
    EMPRESA ABC LTDA, pessoa jurídica de direito privado, já qualificada nos autos do processo em epígrafe, vem, respeitosamente, à presença de Vossa Excelência, por seus advogados que esta subscrevem, apresentar CONTESTAÇÃO à ação proposta por JOÃO DA SILVA, expondo para ao final requerer o seguinte:
    
    DOS FATOS
    
    Trata-se de ação de cobrança proposta pelo autor, alegando ser credor do réu no valor de R$ 10.000,00.
    
    No entanto, tal alegação não merece prosperar, uma vez que o réu já efetuou o pagamento integral dos valores cobrados, conforme comprovantes anexos.
    
    DO DIREITO
    
    Não procedem as alegações do autor, pois os valores já foram pagos, conforme prevê o art. 397 do Código Civil.
    
    A presente ação caracteriza cobrança indevida, devendo ser aplicadas as sanções cabíveis.
    
    DOS PEDIDOS
    
    Ante o exposto, requer-se:
    
    a) A improcedência total dos pedidos do autor;
    b) A condenação do autor ao pagamento das custas processuais e honorários advocatícios;
    c) A produção de todos os meios de prova em direito admitidos.
    
    Termos em que,
    Pede deferimento.
    São Paulo, """ + datetime.datetime.now().strftime("%d/%m/%Y")
    
    # Salvar em arquivo em vez de sessão
    result_id = save_result_to_file(fake_gemini_response)
    session['result_id'] = result_id
    
    logger.info(f"Sessão de teste configurada: {session['test_value']} e result_id: {result_id}")
    
    # Redirecionar para a página de resultado para testar
    return redirect(url_for('resultado', id=result_id))

@app.route('/debug/view_session')
def debug_view_session():
    """Visualizar o conteúdo da sessão atual"""
    session_data = dict(session)
    # Limitar o tamanho da resposta para não sobrecarregar a página
    for key, value in session_data.items():
        if isinstance(value, str) and len(value) > 500:
            session_data[key] = value[:500] + "... [truncado]"
    
    return jsonify({
        'session_id': request.cookies.get('session', 'não encontrado'),
        'session_data': session_data
    })

@app.errorhandler(413)
def request_entity_too_large(error):
    logger.error("Arquivo muito grande enviado")
    return render_template('index.html', error='O arquivo enviado é muito grande. O limite é de 16MB.'), 413

@app.errorhandler(500)
def internal_server_error(error):
    logger.error(f"Erro interno do servidor: {str(error)}")
    return render_template('index.html', error='Erro interno do servidor. Por favor, tente novamente.'), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    logger.info(f"Iniciando servidor na porta {port}")
    app.run(host='0.0.0.0', port=port, debug=True) 